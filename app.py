
import streamlit as st
import cv2
import torch
import torch.nn.functional as F
import numpy as np
import sys, os, io
from PIL import Image, ImageDraw
import torchvision.transforms as T

st.set_page_config(page_title="Handwriting Generator", page_icon="✍️", layout="wide")

@st.cache_resource
def load_models():
    BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
    HIGAN_ROOT = os.path.join(BASE_DIR, "HiGANplus", "HiGAN+")
    BANK_PATH  = os.path.join(BASE_DIR, "glyph_bank.pth")
    sys.path.insert(0, HIGAN_ROOT)
    from munch import munchify
    import yaml
    from networks.module import StyleEncoder, StyleBackbone
    DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
    with open(os.path.join(HIGAN_ROOT, "configs", "gan_iam.yml")) as f:
        cfg = munchify(yaml.safe_load(f))
    cfg.device = DEVICE
    checkpoint = torch.load(os.path.join(HIGAN_ROOT, "pretrained", "deploy_HiGAN+.pth"), map_location=DEVICE)
    style_backbone = StyleBackbone(**cfg.StyBackbone).to(DEVICE)
    style_encoder  = StyleEncoder(**cfg.EncModel).to(DEVICE)
    style_backbone.load_state_dict(checkpoint["StyleBackbone"])
    style_encoder.load_state_dict(checkpoint["StyleEncoder"])
    style_backbone.eval()
    style_encoder.eval()
    sys.path.insert(0, BASE_DIR)
    from glyph_vae import GlyphVAE
    saved = torch.load(BANK_PATH, map_location=DEVICE)
    # ← change latent_dim=128 if you trained 64x64 model
    vae = GlyphVAE(latent_dim=128, style_dim=32, n_chars=36).to(DEVICE)
    vae.load_state_dict(saved["vae_state"])
    vae.char_mu     = saved["char_mu"]
    vae.char_logvar = saved["char_logvar"]
    vae.eval()
    return style_backbone, style_encoder, vae, DEVICE

style_backbone, style_encoder, vae, DEVICE = load_models()

# ── Sidebar ─────────────────────────────────────────────────
st.sidebar.title("✍️ Settings")
page_size    = st.sidebar.selectbox("Page size", ["a4","a5","letter","legal"])
font_size    = st.sidebar.slider("Font size", 20, 60, 32)
line_spacing = st.sidebar.slider("Line spacing", 1.2, 2.5, 1.6)
margin_px    = st.sidebar.slider("Margin", 60, 200, 120)
neatness     = st.sidebar.slider("Neatness", 0.3, 1.5, 0.8)
h1_scale     = st.sidebar.slider("H1 scale", 1.2, 2.5, 1.8)
h2_scale     = st.sidebar.slider("H2 scale", 1.0, 2.0, 1.4)
h1_underline = st.sidebar.toggle("H1 underline", value=True)

# ── Main ────────────────────────────────────────────────────
st.title("✍️ Handwriting Generator")
col1, col2 = st.columns([1, 1])

with col1:
    sample_file  = st.file_uploader("Upload handwriting sample", type=["png","jpg","jpeg"])
    input_text   = st.text_area("Text to convert",
                    "# My Document\n## Introduction\nThe quick brown fox jumps over the lazy dog.",
                    height=200, help="Use # H1, ## H2, ### H3")
    generate_btn = st.button("✍️ Generate", type="primary", use_container_width=True)

with col2:
    st.subheader("Preview")
    preview_placeholder = st.empty()

# ── Helpers ─────────────────────────────────────────────────
CHARSET     = "abcdefghijklmnopqrstuvwxyz0123456789"
CHARSET_SET = set(CHARSET)
CHAR2IDX    = {c: i for i, c in enumerate(CHARSET)}
PAGE_SIZES  = {"a4":(1240,1754),"a5":(874,1240),"letter":(1275,1650),"legal":(1275,2100)}

# ── Style extraction ────────────────────────────────────────
def extract_style(img_bytes):
    img = Image.open(io.BytesIO(img_bytes)).convert("L")

    tf = T.Compose([
        T.Resize((64,256)),
        T.ToTensor(),
        T.Normalize([0.5],[0.5])
    ])

    t  = tf(img).unsqueeze(0).to(DEVICE)
    il = torch.tensor([256]).to(DEVICE)

    with torch.no_grad():
        s = style_encoder(t, il, cnn_backbone=style_backbone, vae_mode=True)

    style_vec = s[1] # checking with mu

    style_vec = F.normalize(style_vec, dim=1) # trying to not normalise so as to not squash
    # st.write(f"style_vec min: {style_vec.min().item():.4f}, max: {style_vec.max().item():.4f}") # debug
    # st.write(f"style_vec: {style_vec[0, :8]}")  # first 8 values

    return style_vec

def tensor_to_pil(t):
    img = (t.squeeze().cpu().numpy() + 1) / 2
    return Image.fromarray((img*255).clip(0,255).astype(np.uint8))

# def generate_glyph(char, style_vec, sigma_mult=1.0):
#     if char not in CHARSET_SET: return None
#     idx  = CHAR2IDX[char]
    
#     oh   = F.one_hot(torch.tensor([idx], device=DEVICE), 36).float()
#     mu   = vae.char_mu[idx].unsqueeze(0).to(DEVICE)
#     lv   = vae.char_logvar[idx].unsqueeze(0).to(DEVICE)
#     std  = torch.exp(0.5 * lv) * sigma_mult
#     noise = torch.randn_like(mu) * 0.8 
#     z    = mu + std * noise    #########################works okay
 
#     sv   = style_vec.reshape(1, -1)
#     cond = torch.cat([sv, oh], dim=1)
    
#     with torch.no_grad():
#         r = vae.decode(z, cond)
#     arr = (r.squeeze().cpu().numpy() + 1) / 2
#     arr = (arr * 255).clip(0, 255).astype(np.uint8)

#     # Sharpen
#     blurred = cv2.GaussianBlur(arr.astype(np.float32), (3,3), 3)
#     sharp   = cv2.addWeighted(arr.astype(np.float32), 1.5, blurred, -0.5, 0)
#     sharp   = np.clip(sharp, 0, 255).astype(np.uint8)

#     # Stretch contrast
#     p5, p95 = np.percentile(sharp, 5), np.percentile(sharp, 95)
#     if p95 > p5:
#         sharp = ((sharp.astype(np.float32) - p5) / (p95 - p5) * 255)
#         sharp = np.clip(sharp, 0, 255).astype(np.uint8)

#     return Image.fromarray(sharp)

def generate_glyph(char, style_vec, sigma_mult=1.0):
    if char not in CHARSET_SET: return None
    idx  = CHAR2IDX[char]
    oh   = F.one_hot(torch.tensor([idx], device=DEVICE), 36).float()
    mu   = vae.char_mu[idx].unsqueeze(0).to(DEVICE)
    lv   = vae.char_logvar[idx].unsqueeze(0).to(DEVICE)
    std  = torch.exp(0.5 * lv) * sigma_mult
    noise = torch.randn_like(mu) * 0.8
    z    = mu + std * noise
    
    sv   = style_vec.reshape(1, -1) * 3.0
    
    # st.write(f"style_vec shape: {sv.shape}, oh shape: {oh.shape}") # debug line
    cond = torch.cat([sv, oh], dim=1)
    # st.write(f"cond first 8: {cond[0, :8]}") #remove later
    with torch.no_grad():
        r = vae.decode(z, cond)
    arr = (r.squeeze().cpu().numpy() + 1) / 2
    arr = (arr * 255).clip(0, 255).astype(np.uint8)
    blurred = cv2.GaussianBlur(arr.astype(np.float32), (3,3), 2)   # (3,3),3
    sharp   = cv2.addWeighted(arr.astype(np.float32), 4.0, blurred, -0.5, 0)   # 1.5->2.5   -0.5->-1.5
    sharp   = np.clip(sharp, 0, 255).astype(np.uint8)
    p5, p95 = np.percentile(sharp, 5), np.percentile(sharp, 95)
    if p95 > p5:
        sharp = ((sharp.astype(np.float32) - p5) / (p95 - p5) * 255)
        sharp = np.clip(sharp, 0, 255).astype(np.uint8)
    # sharp[sharp > 120] = 255  # ← key fix for clarity
    sharp[sharp > 180] = 255 
    return Image.fromarray(sharp)

def parse_text(text):
    lines = []
    for line in text.strip().split("\n"):
        l = line.strip()
        if l.startswith("### "): lines.append(("h3", l[4:]))
        elif l.startswith("## "): lines.append(("h2", l[3:]))
        elif l.startswith("# "): lines.append(("h1", l[2:]))
        elif l: lines.append(("body", l))
        else: lines.append(("blank",""))
    return lines

# def render_pages(text, style_vec, pg_size, base_fs, ls, margin, neatness, h1_sc, h2_sc, h1_ul):
#     HS = {
#         "h1":  {"scale":h1_sc,  "sigma":0.3, "underline":h1_ul, "dilate":True},
#         "h2":  {"scale":h2_sc,  "sigma":0.5, "underline":False,  "dilate":False},
#         "h3":  {"scale":1.15,   "sigma":0.7, "underline":False,  "dilate":False},
#         "body":{"scale":1.0,    "sigma":neatness,"underline":False,"dilate":False},
#     }
#     PW, PH = PAGE_SIZES[pg_size]
#     pages  = []
#     page   = Image.new("L",(PW,PH),255)
#     draw   = ImageDraw.Draw(page)
#     x, y   = margin, margin

#     def new_page():
#         nonlocal page, draw, x, y
#         pages.append(page)
#         page = Image.new("L",(PW,PH),255)
#         draw = ImageDraw.Draw(page)
#         x, y = margin, margin

#     for level, txt in parse_text(text):
#         if level == "blank": y += base_fs; continue
#         hs     = HS[level]
#         ch     = int(base_fs * hs["scale"])
#         cw     = int(ch * 0.75)
#         sw     = int(cw * 0.5)
#         sigma  = hs["sigma"]
#         y += int(ch * (0.3 if level!="body" else 0))
#         if y + ch > PH - margin: new_page()

#         words = txt.split(); cur = []; cx = 0
#         wrapped = []
#         for w in words:
#             ww = len(w)*cw + sw
#             if cx+ww > PW-2*margin and cur:
#                 wrapped.append(cur); cur=[w]; cx=ww
#             else:
#                 cur.append(w); cx+=ww
#         if cur: wrapped.append(cur)

#         for li, wline in enumerate(wrapped):
#             if y+ch > PH-margin: new_page()
#             x = margin; ax = margin
#             for word in wline:
#                 wc = "".join(c for c in word.lower() if c in CHARSET_SET)
#                 for char in wc:
#                     if x+cw > PW-margin: break
#                     g = generate_glyph(char, style_vec, sigma)
#                     if g is None: x+=cw; continue
#                     g = g.resize((cw,ch), Image.LANCZOS)
#                     if hs["dilate"]:
#                         a = np.array(g)
#                         a = cv2.dilate(255-a, np.ones((2,2),np.uint8))
#                         g = Image.fromarray(255-a)
#                     jit = np.random.randint(-2,3)
#                     ga  = np.array(g)
#                     page.paste(Image.fromarray(np.zeros_like(ga)),
#                                (x, y+jit),
#                                mask=Image.fromarray((255-ga).astype(np.uint8)))
#                     x += cw; ax = x
#                 x += sw
#             if hs["underline"] and li==len(wrapped)-1:
#                 draw.line([(margin,y+ch+4),(ax,y+ch+4)],fill=0,width=2)
#             y += int(ch*ls)
#         y += int(ch*0.2)

#     pages.append(page)
#     return pages

def render_pages(text, style_vec, pg_size, base_fs, ls, margin, neatness, h1_sc, h2_sc, h1_ul):
    HS = {
        "h1":  {"scale":h1_sc,  "sigma":0.3, "underline":h1_ul, "dilate":True},
        "h2":  {"scale":h2_sc,  "sigma":0.5, "underline":False,  "dilate":False},
        "h3":  {"scale":1.15,   "sigma":0.7, "underline":False,  "dilate":False},
        "body":{"scale":1.0,    "sigma":neatness,"underline":False,"dilate":False},
    }
    PW, PH = PAGE_SIZES[pg_size]
    pages  = []
    page   = Image.new("L",(PW,PH),255)
    draw   = ImageDraw.Draw(page)
    x, y   = margin, margin

    def new_page():
        nonlocal page, draw, x, y
        pages.append(page)
        page = Image.new("L",(PW,PH),255)
        draw = ImageDraw.Draw(page)
        x, y = margin, margin

    for level, txt in parse_text(text):
        if level == "blank": y += base_fs; continue
        hs  = HS[level]
        # Fixed char size like Kaggle — more consistent output
        ch    = int(base_fs * hs["scale"])  # scale with font_size slider
        cw    = int(ch * 0.75)
        sw  = int(cw * 0.5)
        sigma = hs["sigma"]
        y += int(base_fs * hs["scale"] * (0.3 if level != "body" else 0))
        if y + ch > PH - margin: new_page()

        words = txt.split()
        for word in words:
            wc = "".join(c for c in word.lower() if c in CHARSET_SET)
            ww = len(wc) * cw
            if x + ww > PW - margin and x > margin:
                x  = margin
                y += int(ch * 1.4)
            if y + ch > PH - margin: new_page()
            for char in wc:
                if x + cw > PW - margin: break
                g = generate_glyph(char, style_vec, sigma)
                if g is None: x += cw; continue
                g = g.resize((cw, ch), Image.LANCZOS)
                if hs["dilate"]:
                    a = np.array(g)
                    a = cv2.dilate(255-a, np.ones((2,2),np.uint8))
                    g = Image.fromarray(255-a)
                jit = np.random.randint(-2, 3)
                ga  = np.array(g)
                page.paste(Image.fromarray(np.zeros_like(ga)),
                           (x, y+jit),
                           mask=Image.fromarray((255-ga).astype(np.uint8)))
                x += cw
            if hs["underline"]:
                draw.line([(margin, y+ch+3), (x, y+ch+3)], fill=0, width=2)
            x += sw + np.random.randint(-2, 3)  # random space jitter like Kaggle
        x  = margin
        y += int(ch * 1.6)

    pages.append(page)
    return pages

# ── Generate ─────────────────────────────────────────────────
if generate_btn:
    if not sample_file:
        st.warning("Upload a handwriting sample first.")
    else:
        with st.spinner("Extracting style..."):
            style_vec = extract_style(sample_file.read())
        with st.spinner("Rendering..."):
            rendered = render_pages(input_text, style_vec, page_size, font_size,
                                    line_spacing, margin_px, neatness,
                                    h1_scale, h2_scale, h1_underline)
        preview_placeholder.image(rendered[0], caption="Preview", use_column_width=True)

        st.markdown("---")
        dl1, dl2 = st.columns(2)

        # PDF
        from reportlab.lib.pagesizes import A4,A5,LETTER,LEGAL
        from reportlab.pdfgen import canvas as pc
        rl = {"a4":A4,"a5":A5,"letter":LETTER,"legal":LEGAL}
        pdf_buf = io.BytesIO()
        c = pc.Canvas(pdf_buf, pagesize=rl[page_size])
        pw, ph = rl[page_size]
        for pg in rendered:
            b = io.BytesIO()
            pg.resize((int(pw),int(ph)),Image.LANCZOS).save(b,format="PNG")
            b.seek(0)
            c.drawImage(pc.ImageReader(b),0,0,width=pw,height=ph)
            c.showPage()
        c.save(); pdf_buf.seek(0)
        with dl1:
            st.download_button("⬇️ Download PDF", pdf_buf,
                               "handwriting.pdf","application/pdf",
                               use_container_width=True)

        # DOCX
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        docx_buf = io.BytesIO()
        doc = Document()
        wi = {"a4":8.27,"a5":5.83,"letter":8.5,"legal":8.5}
        for i, pg in enumerate(rendered):
            if i > 0: doc.add_page_break()
            b = io.BytesIO(); pg.save(b,format="PNG"); b.seek(0)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(b, width=Inches(wi[page_size]))
        doc.save(docx_buf); docx_buf.seek(0)
        with dl2:
            st.download_button("⬇️ Download DOCX", docx_buf,
                               "handwriting.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True)

        if len(rendered) > 1:
            st.info(f"{len(rendered)} pages — all included in download.")
